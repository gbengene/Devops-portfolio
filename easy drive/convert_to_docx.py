"""Convert Easy Drive markdown files to Word (.docx) documents."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BASE_DIR = Path(r"C:\projects\easy drive")

FILES = [
    ("BUSINESS_PLAN.md",    "Easy Drive — Business Plan.docx"),
    ("ARCHITECTURE.md",     "Easy Drive — Architecture.docx"),
    ("PROJECT_PLAN.md",     "Easy Drive — Project Plan.docx"),
    ("FINANCIAL_GUIDANCE.md", "Easy Drive — Financial Guidance.docx"),
]

# Brand colours
BRAND_BLUE  = RGBColor(0x1A, 0x56, 0xDB)   # heading blue
BRAND_DARK  = RGBColor(0x11, 0x18, 0x27)   # near-black body
TABLE_HEADER_BG = "1A56DB"
TABLE_ROW_ALT   = "EFF6FF"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **borders):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in borders.items():
        el = OxmlElement(f"w:{edge}")
        for k, v in attrs.items():
            el.set(qn(f"w:{k}"), v)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def style_heading(paragraph, level: int):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    run.font.color.rgb = BRAND_BLUE
    run.font.bold = True
    sizes = {1: 22, 2: 16, 3: 13, 4: 12}
    run.font.size = Pt(sizes.get(level, 12))
    if level == 1:
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after  = Pt(4)
    else:
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after  = Pt(2)


def add_cover(doc: Document, title: str, subtitle: str):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Easy Drive")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = BRAND_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.size = Pt(20)
    run2.font.bold = True
    run2.font.color.rgb = BRAND_DARK

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(subtitle)
    run3.font.size = Pt(11)
    run3.font.italic = True
    run3.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_page_break()


def add_table(doc: Document, rows: list[list[str]]):
    """Render a markdown table. rows[0] is the header row."""
    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = cell_text.strip()
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run = p.runs[0] if p.runs else p.add_run(cell_text.strip())

            if r_idx == 0:
                # Header row
                set_cell_bg(cell, TABLE_HEADER_BG)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(10)
            else:
                bg = TABLE_ROW_ALT if r_idx % 2 == 0 else "FFFFFF"
                set_cell_bg(cell, bg)
                run.font.size = Pt(10)
                run.font.color.rgb = BRAND_DARK

    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------

def parse_table(lines: list[str]) -> list[list[str]] | None:
    """Return parsed rows if lines look like a markdown table, else None."""
    clean = [l for l in lines if l.strip() and not re.match(r"^\|[-| :]+\|$", l.strip())]
    if not clean:
        return None
    rows = []
    for line in clean:
        line = line.strip()
        if line.startswith("|"):
            line = line[1:]
        if line.endswith("|"):
            line = line[:-1]
        rows.append([c.strip() for c in line.split("|")])
    return rows if rows else None


def render_inline(text: str) -> tuple[str, bool, bool]:
    """Strip inline markdown, return (text, bold, italic)."""
    bold = False
    italic = False
    t = text
    # Bold+italic
    t, n = re.subn(r"\*\*\*(.+?)\*\*\*", r"\1", t)
    if n:
        bold = italic = True
    # Bold
    t, n = re.subn(r"\*\*(.+?)\*\*", r"\1", t)
    if n:
        bold = True
    # Italic
    t, n = re.subn(r"\*(.+?)\*|_(.+?)_", lambda m: m.group(1) or m.group(2), t)
    if n:
        italic = True
    # Inline code
    t = re.sub(r"`(.+?)`", r"\1", t)
    # Links
    t = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", t)
    return t, bold, italic


def convert_md_to_docx(md_path: Path, docx_path: Path, title: str, subtitle: str):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # Default body style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BRAND_DARK

    add_cover(doc, title, subtitle)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    i = 0
    table_buffer: list[str] = []
    code_buffer:  list[str] = []
    in_code = False

    def flush_table():
        if table_buffer:
            rows = parse_table(table_buffer)
            if rows:
                add_table(doc, rows)
            table_buffer.clear()

    def flush_code():
        if code_buffer:
            add_code_block(doc, code_buffer)
            code_buffer.clear()

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # Fenced code block toggle
        if stripped.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(raw)
            i += 1
            continue

        # Table row
        if stripped.startswith("|"):
            flush_table()  # flush previous table first (shouldn't be needed)
            table_buffer.append(stripped)
            i += 1
            continue
        else:
            flush_table()

        # Separator row (---) — skip
        if re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            text, bold, italic = render_inline(m.group(2))
            heading_styles = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
            p = doc.add_paragraph(style=heading_styles.get(level, "Heading 3"))
            p.clear()
            run = p.add_run(text)
            run.font.bold = True
            run.font.color.rgb = BRAND_BLUE
            sizes = {1: 22, 2: 16, 3: 13, 4: 12}
            run.font.size = Pt(sizes.get(level, 12))
            if level == 1:
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after  = Pt(6)
            elif level == 2:
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after  = Pt(4)
            else:
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after  = Pt(2)
            i += 1
            continue

        # Bullet / list item
        m = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)", raw)
        if m:
            indent_level = len(m.group(1)) // 2
            text, bold, italic = render_inline(m.group(3))
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.bold = bold
            run.font.italic = italic
            run.font.color.rgb = BRAND_DARK
            i += 1
            continue

        # Checkbox list item
        m = re.match(r"^(\s*)[-*]\s+\[[ x]\]\s+(.*)", raw)
        if m:
            text, bold, italic = render_inline(m.group(2))
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run("☐  " + text)
            run.font.size = Pt(11)
            run.font.color.rgb = BRAND_DARK
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        text, bold, italic = render_inline(stripped)
        if not text:
            i += 1
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = BRAND_DARK
        i += 1

    flush_table()
    flush_code()

    doc.save(docx_path)
    print(f"  Saved: {docx_path.name}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

SUBTITLES = {
    "BUSINESS_PLAN.md":     "Marketplace Model — Version 2.0 | May 2026",
    "ARCHITECTURE.md":      "Solution Architecture — Version 2.0 | May 2026",
    "PROJECT_PLAN.md":      "Launch Project Plan | May 2026",
    "FINANCIAL_GUIDANCE.md":"Financial & Accounting Guidance | May 2026",
}

TITLES = {
    "BUSINESS_PLAN.md":     "Business Plan",
    "ARCHITECTURE.md":      "Solution Architecture",
    "PROJECT_PLAN.md":      "Project Plan",
    "FINANCIAL_GUIDANCE.md":"Financial Guidance",
}

if __name__ == "__main__":
    print("Converting Easy Drive documents to Word format...\n")
    for md_name, docx_name in FILES:
        md_path   = BASE_DIR / md_name
        docx_path = BASE_DIR / docx_name
        if not md_path.exists():
            print(f"  SKIP (not found): {md_name}")
            continue
        print(f"  Converting: {md_name}")
        convert_md_to_docx(
            md_path, docx_path,
            title=TITLES[md_name],
            subtitle=SUBTITLES[md_name],
        )
    print("\nDone. All Word documents saved to:")
    print(f"  {BASE_DIR}")
